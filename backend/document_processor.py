"""
Medical Document Processing Module

This module handles processing and ingestion of medical documents
into the vector database for RAG functionality.
"""

import os
import asyncio
from typing import List, Dict, Any, Optional
from pathlib import Path
import json
import structlog
from dataclasses import dataclass
import pandas as pd
from datetime import datetime

# Document processing imports
import PyPDF2
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import requests

from vector_store import PineconeVectorStore, Document

logger = structlog.get_logger()

@dataclass
class MedicalDocument:
    """Medical document with metadata"""
    title: str
    content: str
    source: str
    document_type: str
    specialty: Optional[str] = None
    date_published: Optional[datetime] = None
    authors: Optional[List[str]] = None
    keywords: Optional[List[str]] = None
    credibility_score: float = 1.0

class MedicalDocumentProcessor:
    """
    Processes various types of medical documents for RAG ingestion
    """
    
    def __init__(self, vector_store: PineconeVectorStore):
        self.vector_store = vector_store
        self.supported_formats = ['.pdf', '.txt', '.docx', '.json', '.csv']
        self.medical_specialties = [
            'cardiology', 'dermatology', 'endocrinology', 'gastroenterology',
            'hematology', 'immunology', 'infectious_disease', 'nephrology',
            'neurology', 'oncology', 'pediatrics', 'psychiatry', 'pulmonology',
            'rheumatology', 'surgery', 'urology', 'gynecology', 'orthopedics',
            'ophthalmology', 'otolaryngology', 'radiology', 'pathology',
            'anesthesiology', 'emergency_medicine', 'family_medicine',
            'internal_medicine', 'preventive_medicine', 'geriatrics'
        ]
    
    async def process_directory(self, directory_path: str) -> bool:
        """Process all medical documents in a directory"""
        try:
            logger.info(f"Processing medical documents from: {directory_path}")
            
            directory = Path(directory_path)
            if not directory.exists():
                raise FileNotFoundError(f"Directory not found: {directory_path}")
            
            documents = []
            processed_count = 0
            
            # Process all supported files
            for file_path in directory.rglob("*"):
                if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                    try:
                        doc = await self.process_file(str(file_path))
                        if doc:
                            documents.append(doc)
                            processed_count += 1
                            logger.info(f"Processed: {file_path.name}")
                    except Exception as e:
                        logger.error(f"Error processing {file_path}: {str(e)}")
            
            # Convert to vector store documents and add to database
            if documents:
                vector_documents = self._convert_to_vector_documents(documents)
                await self.vector_store.add_documents(vector_documents)
                logger.info(f"Successfully processed and indexed {processed_count} medical documents")
            
            return True
            
        except Exception as e:
            logger.error(f"Error processing directory: {str(e)}")
            raise
    
    async def process_file(self, file_path: str) -> Optional[MedicalDocument]:
        """Process a single medical document file"""
        try:
            path = Path(file_path)
            suffix = path.suffix.lower()
            
            if suffix == '.pdf':
                return await self._process_pdf(file_path)
            elif suffix == '.txt':
                return await self._process_text(file_path)
            elif suffix == '.docx':
                return await self._process_docx(file_path)
            elif suffix == '.json':
                return await self._process_json(file_path)
            elif suffix == '.csv':
                return await self._process_csv(file_path)
            else:
                logger.warning(f"Unsupported file format: {suffix}")
                return None
                
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            return None
    
    async def _process_pdf(self, file_path: str) -> Optional[MedicalDocument]:
        """Process PDF medical document"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                metadata = pdf_reader.metadata or {}
                title = metadata.get('/Title', Path(file_path).stem)
                authors = metadata.get('/Author', '').split(';') if metadata.get('/Author') else []
                
                # Extract text content
                content_parts = []
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        content_parts.append(text.strip())
                
                content = '\n\n'.join(content_parts)
                
                if not content.strip():
                    logger.warning(f"No extractable text found in PDF: {file_path}")
                    return None
                
                # Determine specialty from content
                specialty = self._detect_medical_specialty(content, title)
                
                return MedicalDocument(
                    title=title,
                    content=content,
                    source=file_path,
                    document_type='pdf',
                    specialty=specialty,
                    authors=authors if authors else None,
                    keywords=self._extract_keywords(content),
                    credibility_score=self._assess_credibility(content, title, authors)
                )
                
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            return None
    
    async def _process_text(self, file_path: str) -> Optional[MedicalDocument]:
        """Process plain text medical document"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            if not content.strip():
                return None
            
            title = Path(file_path).stem
            specialty = self._detect_medical_specialty(content, title)
            
            return MedicalDocument(
                title=title,
                content=content,
                source=file_path,
                document_type='text',
                specialty=specialty,
                keywords=self._extract_keywords(content),
                credibility_score=self._assess_credibility(content, title)
            )
            
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {str(e)}")
            return None
    
    async def _process_docx(self, file_path: str) -> Optional[MedicalDocument]:
        """Process Word document"""
        try:
            doc = DocxDocument(file_path)
            
            # Extract content
            content_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text.strip())
            
            content = '\n\n'.join(content_parts)
            
            if not content.strip():
                return None
            
            # Extract metadata
            core_props = doc.core_properties
            title = core_props.title or Path(file_path).stem
            authors = [core_props.author] if core_props.author else []
            
            specialty = self._detect_medical_specialty(content, title)
            
            return MedicalDocument(
                title=title,
                content=content,
                source=file_path,
                document_type='docx',
                specialty=specialty,
                authors=authors,
                keywords=self._extract_keywords(content),
                credibility_score=self._assess_credibility(content, title, authors)
            )
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            return None
    
    async def _process_json(self, file_path: str) -> Optional[MedicalDocument]:
        """Process JSON medical data"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Handle different JSON structures
            if isinstance(data, dict):
                if 'title' in data and 'content' in data:
                    # Single document format
                    return self._create_document_from_json(data, file_path)
                elif 'documents' in data:
                    # Multiple documents format - process first one for now
                    if data['documents']:
                        return self._create_document_from_json(data['documents'][0], file_path)
            elif isinstance(data, list):
                # Array of documents - process first one
                if data:
                    return self._create_document_from_json(data[0], file_path)
            
            return None
            
        except Exception as e:
            logger.error(f"Error processing JSON {file_path}: {str(e)}")
            return None
    
    def _create_document_from_json(self, data: Dict[str, Any], file_path: str) -> MedicalDocument:
        """Create MedicalDocument from JSON data"""
        title = data.get('title', Path(file_path).stem)
        content = data.get('content', data.get('text', ''))
        
        return MedicalDocument(
            title=title,
            content=content,
            source=file_path,
            document_type='json',
            specialty=data.get('specialty') or self._detect_medical_specialty(content, title),
            authors=data.get('authors'),
            keywords=data.get('keywords') or self._extract_keywords(content),
            credibility_score=data.get('credibility_score', 1.0)
        )
    
    async def _process_csv(self, file_path: str) -> Optional[MedicalDocument]:
        """Process CSV medical data"""
        try:
            df = pd.read_csv(file_path)
            
            # Look for content columns
            content_columns = ['content', 'text', 'description', 'abstract', 'summary']
            title_columns = ['title', 'name', 'subject', 'topic']
            
            content_col = None
            title_col = None
            
            for col in content_columns:
                if col in df.columns:
                    content_col = col
                    break
            
            for col in title_columns:
                if col in df.columns:
                    title_col = col
                    break
            
            if not content_col:
                logger.warning(f"No content column found in CSV: {file_path}")
                return None
            
            # Combine all rows into a single document
            content_parts = []
            for _, row in df.iterrows():
                if pd.notna(row[content_col]):
                    title_part = f"**{row[title_col]}**\n" if title_col and pd.notna(row[title_col]) else ""
                    content_parts.append(f"{title_part}{row[content_col]}")
            
            content = '\n\n'.join(content_parts)
            title = Path(file_path).stem
            
            return MedicalDocument(
                title=title,
                content=content,
                source=file_path,
                document_type='csv',
                specialty=self._detect_medical_specialty(content, title),
                keywords=self._extract_keywords(content),
                credibility_score=1.0
            )
            
        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {str(e)}")
            return None
    
    def _detect_medical_specialty(self, content: str, title: str) -> Optional[str]:
        """Detect medical specialty from content and title"""
        text = (content + " " + title).lower()
        
        specialty_keywords = {
            'cardiology': ['heart', 'cardiac', 'cardiovascular', 'coronary', 'arrhythmia', 'ecg', 'ekg'],
            'dermatology': ['skin', 'dermatitis', 'eczema', 'psoriasis', 'melanoma', 'rash'],
            'endocrinology': ['diabetes', 'thyroid', 'hormone', 'insulin', 'glucose', 'endocrine'],
            'gastroenterology': ['stomach', 'intestine', 'digestive', 'gastric', 'colon', 'bowel'],
            'neurology': ['brain', 'neurological', 'seizure', 'stroke', 'alzheimer', 'parkinson'],
            'oncology': ['cancer', 'tumor', 'malignant', 'chemotherapy', 'radiation', 'oncology'],
            'pediatrics': ['child', 'children', 'pediatric', 'infant', 'baby', 'adolescent'],
            'psychiatry': ['mental', 'depression', 'anxiety', 'psychiatric', 'psychology', 'therapy'],
            'pulmonology': ['lung', 'respiratory', 'breathing', 'asthma', 'copd', 'pneumonia'],
            'orthopedics': ['bone', 'joint', 'fracture', 'orthopedic', 'musculoskeletal', 'spine']
        }
        
        for specialty, keywords in specialty_keywords.items():
            if any(keyword in text for keyword in keywords):
                return specialty
        
        return None
    
    def _extract_keywords(self, content: str) -> List[str]:
        """Extract medical keywords from content"""
        # Simple keyword extraction - could be enhanced with NLP
        medical_terms = [
            'diagnosis', 'treatment', 'symptoms', 'patient', 'disease', 'condition',
            'medication', 'therapy', 'clinical', 'medical', 'health', 'syndrome',
            'disorder', 'infection', 'chronic', 'acute', 'prevention', 'screening'
        ]
        
        content_lower = content.lower()
        found_keywords = [term for term in medical_terms if term in content_lower]
        
        return found_keywords[:10]  # Limit to top 10
    
    def _assess_credibility(
        self, 
        content: str, 
        title: str, 
        authors: Optional[List[str]] = None
    ) -> float:
        """Assess document credibility score (0.0 to 1.0)"""
        score = 0.5  # Base score
        
        # Check for academic indicators
        academic_indicators = [
            'study', 'research', 'clinical trial', 'peer-reviewed',
            'journal', 'university', 'hospital', 'doi:', 'pmid:'
        ]
        
        text = (content + " " + title).lower()
        
        for indicator in academic_indicators:
            if indicator in text:
                score += 0.1
        
        # Check for author credentials
        if authors:
            for author in authors:
                if any(title in author.lower() for title in ['dr.', 'md', 'phd', 'prof']):
                    score += 0.1
                    break
        
        # Check content length (longer content often more comprehensive)
        if len(content) > 5000:
            score += 0.1
        elif len(content) > 2000:
            score += 0.05
        
        return min(score, 1.0)  # Cap at 1.0
    
    def _convert_to_vector_documents(self, medical_docs: List[MedicalDocument]) -> List[Document]:
        """Convert MedicalDocument objects to vector store Document objects"""
        vector_docs = []
        
        for doc in medical_docs:
            # Chunk long documents
            chunks = self.vector_store.chunk_text(doc.content)
            
            for i, chunk in enumerate(chunks):
                metadata = {
                    'title': doc.title,
                    'source': doc.source,
                    'document_type': doc.document_type,
                    'specialty': doc.specialty,
                    'chunk_index': i,
                    'total_chunks': len(chunks),
                    'credibility_score': doc.credibility_score,
                    'date_indexed': datetime.now().isoformat()
                }
                
                if doc.authors:
                    metadata['authors'] = ', '.join(doc.authors)
                
                if doc.keywords:
                    metadata['keywords'] = ', '.join(doc.keywords)
                
                vector_docs.append(Document(
                    content=chunk,
                    metadata=metadata,
                    id=f"{Path(doc.source).stem}_chunk_{i}"
                ))
        
        return vector_docs
    
    async def add_medical_knowledge_samples(self):
        """Add sample medical knowledge for testing"""
        try:
            logger.info("Adding sample medical knowledge...")
            
            sample_docs = [
                MedicalDocument(
                    title="Hypertension Management Guidelines",
                    content="""
                    Hypertension, or high blood pressure, is a common cardiovascular condition affecting millions worldwide. 
                    Normal blood pressure is typically below 120/80 mmHg. Hypertension is diagnosed when blood pressure 
                    consistently measures 140/90 mmHg or higher.
                    
                    Treatment approaches include:
                    1. Lifestyle modifications: Regular exercise, healthy diet (DASH diet), sodium restriction, weight management
                    2. Medications: ACE inhibitors, ARBs, calcium channel blockers, diuretics
                    3. Regular monitoring and follow-up
                    
                    Complications of untreated hypertension include stroke, heart attack, kidney disease, and heart failure.
                    Early detection and proper management are crucial for preventing these serious complications.
                    """,
                    source="sample_hypertension_guide.txt",
                    document_type="text",
                    specialty="cardiology",
                    keywords=["hypertension", "blood pressure", "cardiovascular", "treatment"],
                    credibility_score=0.9
                ),
                
                MedicalDocument(
                    title="Type 2 Diabetes Management",
                    content="""
                    Type 2 diabetes is a chronic metabolic disorder characterized by insulin resistance and relative insulin deficiency.
                    It affects how the body processes glucose, leading to elevated blood sugar levels.
                    
                    Risk factors include:
                    - Obesity and sedentary lifestyle
                    - Family history of diabetes
                    - Age over 45
                    - Certain ethnicities
                    
                    Management strategies:
                    1. Blood glucose monitoring
                    2. Dietary modifications (carbohydrate counting, portion control)
                    3. Regular physical activity
                    4. Medications: Metformin, insulin, other antidiabetic drugs
                    5. Regular screening for complications
                    
                    Complications can include diabetic retinopathy, nephropathy, neuropathy, and increased cardiovascular risk.
                    Good glycemic control (HbA1c < 7%) significantly reduces complication risk.
                    """,
                    source="sample_diabetes_guide.txt",
                    document_type="text",
                    specialty="endocrinology",
                    keywords=["diabetes", "glucose", "insulin", "blood sugar"],
                    credibility_score=0.9
                ),
                
                MedicalDocument(
                    title="Common Cold vs Flu Symptoms",
                    content="""
                    Both common cold and influenza are respiratory illnesses, but they are caused by different viruses
                    and have distinct symptom patterns.
                    
                    Common Cold symptoms:
                    - Gradual onset
                    - Runny or stuffy nose
                    - Sneezing
                    - Mild body aches
                    - Low-grade fever (rare)
                    - Duration: 7-10 days
                    
                    Influenza (Flu) symptoms:
                    - Sudden onset
                    - High fever (100-104°F)
                    - Severe body aches
                    - Fatigue and weakness
                    - Dry cough
                    - Headache
                    - Duration: 1-2 weeks
                    
                    Treatment:
                    Cold: Rest, fluids, symptom relief
                    Flu: Antiviral medications (if started early), rest, fluids
                    
                    Seek medical attention for severe symptoms, high fever, or difficulty breathing.
                    """,
                    source="sample_cold_flu_guide.txt",
                    document_type="text",
                    specialty="family_medicine",
                    keywords=["cold", "flu", "respiratory", "symptoms"],
                    credibility_score=0.8
                )
            ]
            
            vector_docs = self._convert_to_vector_documents(sample_docs)
            await self.vector_store.add_documents(vector_docs)
            
            logger.info(f"Added {len(sample_docs)} sample medical documents")
            return True
            
        except Exception as e:
            logger.error(f"Error adding sample medical knowledge: {str(e)}")
            raise
